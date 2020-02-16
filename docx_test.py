from docx import Document
import os
import re
import sys

def change_to_mw(line):
    ''' 将数据换算为毫瓦 '''
    ll = line.split()
    newline=''
    for item in ll[:-1]:
        item = round( float(item)*1000 , 6)
        print(str(item))
        newline += str(item) + ' ' 
    return newline

def add_row(data):
    row_new = table.add_row()
    ll = data.split()
    print(ll)
    i=0
    for cell in row_new.cells:
        cell.text =  ll[i]
        #print(ll[i])
        #cell.text = re.match(r'(.+)e-(\d+)',ll[i]).group(1)
        i +=1

if __name__ == '__main__':
    doc = Document()

    doc.add_heading('power analysis',0)
    table = doc.add_table(rows=1,cols=5,style="Table Grid")
    header = table.rows[0].cells
    header[0].text = 'Heir'
    header[1].text = 'Internal (mW)'
    header[2].text = 'Switch (mW)'
    header[3].text = 'Leak (mW)'
    header[4].text = 'Total (mW)'


    try:
        if sys.argv[1] == 'pka_log':
            read_f=open("pka_result.log",'w')
            mod_f=open("pkamodules.txt",'r')
        elif sys.argv[1] == 'spacc_log':
            read_f=open("spacc_result.log",'w')
            mod_f=open("spacc_modules.txt",'r')
        elif sys.argv[1] == 'otg_log':
            read_f=open("otg_result.log",'w')
            mod_f=open("otgmodules.txt",'r')
    except:
        read_f=open("pka_idle.log",'r')
        mod_f=open("pkamodules.txt",'r')

    str_item = ''
    mod_f.seek(0)

    for item in mod_f.readlines():
        str_item = '\s*' + item.strip() 
        read_f.seek(0)
        while True:
            line = read_f.readline()
            if line:
                if re.match(str_item,line) :
                    dataline = read_f.readline()
                    while dataline.strip() == '':
                         dataline = read_f.readline()
                    data = item.split()[0] + ' '+ change_to_mw(dataline)
                    add_row(data)
                    break
            else:
                break

    doc.save('./doc_power_result.docx')
